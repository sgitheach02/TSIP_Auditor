"""Génère le gabarit Word `templates/audit_report_template.docx`.

Ce script est un outil de build, exécuté une seule fois (ou après toute
évolution de la charte graphique) pour produire le fichier `.docx` binaire
contenant les balises Jinja2 (`docxtpl`) consommées par `core.reporter`. Le
gabarit généré est ensuite versionné comme un artefact binaire du projet, au
même titre qu'un logo ou une charte graphique fournie par un graphiste.

Usage : python -m ivsa.scripts.build_report_template
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

_ILEXIA_BLUE = RGBColor(0x0B, 0x3D, 0x91)
_ILEXIA_GREY = RGBColor(0x59, 0x59, 0x59)

TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "templates" / "audit_report_template.docx"


def _set_cell_shading(cell, hex_color: str) -> None:
    shading_element = cell._tc.get_or_add_tcPr()
    shd = shading_element.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_color})
    shading_element.append(shd)


def _add_tag_paragraph(document: Document, tag: str) -> None:
    """Ajoute un paragraphe contenant une unique balise Jinja, dans un run
    unique et non stylé (indispensable pour que docxtpl détecte la balise :
    une balise fragmentée sur plusieurs `runs` XML n'est pas reconnue)."""

    paragraph = document.add_paragraph()
    paragraph.add_run(tag)


def _add_tagged_run(paragraph, tag: str, bold: bool = False, size: int | None = None):
    run = paragraph.add_run(tag)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return run


def build() -> Path:
    document = Document()

    # ------------------------------------------------------------------ Styles
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)

    for level, size in ((1, 20), (2, 15), (3, 12.5)):
        heading_style = document.styles[f"Heading {level}"]
        heading_style.font.color.rgb = _ILEXIA_BLUE
        heading_style.font.size = Pt(size)
        heading_style.font.bold = True

    # ------------------------------------------------------------------ Page de titre
    title = document.add_heading("Rapport d'Audit de Conformité et de Sécurité", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_tagged_run(subtitle, "Trunk SIP ToIP — {{ metadata.client_name }}", size=14)

    generated_by = document.add_paragraph()
    generated_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = generated_by.add_run(
        "Généré automatiquement par IVSA (Ilexia Validation & Security Automator)"
    )
    run.italic = True
    run.font.color.rgb = _ILEXIA_GREY

    document.add_paragraph()

    # ------------------------------------------------------------------ Informations générales
    document.add_heading("1. Informations générales", level=1)
    info_table = document.add_table(rows=0, cols=2)
    info_table.style = "Light Grid Accent 1"
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_rows = [
        ("Client", "{{ metadata.client_name }}"),
        ("Site", "{{ metadata.site_name }}"),
        ("Date d'analyse", "{{ metadata.analysis_timestamp }}"),
        ("Analyste", "{{ metadata.analyst }}"),
        ("Fichier de capture analysé", "{{ metadata.pcap_filename }}"),
        ("Éditeur / Version IPBX", "{{ metadata.ipbx_vendor }} {{ metadata.ipbx_version }}"),
        ("Opérateur", "{{ metadata.operator }}"),
        ("Type de Trunk", "{{ metadata.trunk_type }}"),
        ("Trames capturées / Messages SIP / Dialogues", "{{ metadata.total_packets_captured }} / {{ metadata.total_sip_messages }} / {{ metadata.total_sip_dialogs }}"),
        ("Référentiel de règles appliqué", "{{ metadata.ruleset_name }} (v{{ metadata.ruleset_version }})"),
    ]
    for label, tag in info_rows:
        row_cells = info_table.add_row().cells
        row_cells[0].paragraphs[0].add_run(label).bold = True
        row_cells[1].paragraphs[0].add_run(tag)

    document.add_paragraph()

    # ------------------------------------------------------------------ Synthèse
    document.add_heading("2. Synthèse des résultats", level=1)

    summary_table = document.add_table(rows=1, cols=2)
    summary_table.style = "Light Grid Accent 1"
    header_cells = summary_table.rows[0].cells
    header_cells[0].paragraphs[0].add_run("Statut du point de contrôle").bold = True
    header_cells[1].paragraphs[0].add_run("Nombre").bold = True
    for hex_color, cells in ((("D9E2F3"), header_cells),):
        for cell in cells:
            _set_cell_shading(cell, hex_color)

    for label, tag in (
        ("Conforme (PASS)", "{{ pass_count }}"),
        ("Non conforme (FAIL)", "{{ fail_count }}"),
        ("Avertissement (WARNING)", "{{ warning_count }}"),
    ):
        row_cells = summary_table.add_row().cells
        row_cells[0].paragraphs[0].add_run(label)
        row_cells[1].paragraphs[0].add_run(tag)

    document.add_paragraph()
    document.add_paragraph().add_run("Répartition des anomalies par sévérité :").bold = True

    severity_table = document.add_table(rows=1, cols=2)
    severity_table.style = "Light Grid Accent 1"
    sev_header = severity_table.rows[0].cells
    sev_header[0].paragraphs[0].add_run("Sévérité").bold = True
    sev_header[1].paragraphs[0].add_run("Nombre d'anomalies").bold = True
    for cell in sev_header:
        _set_cell_shading(cell, "D9E2F3")

    sev_open_marker = severity_table.add_row().cells
    sev_open_marker[0].paragraphs[0].add_run("{%tr for sev in severity_summary %}")

    sev_row = severity_table.add_row().cells
    sev_row[0].paragraphs[0].add_run("{{ sev.label }}")
    sev_row[1].paragraphs[0].add_run("{{ sev.count }}")

    sev_close_marker = severity_table.add_row().cells
    sev_close_marker[0].paragraphs[0].add_run("{%tr endfor %}")

    document.add_paragraph()

    # ------------------------------------------------------------------ Détail des points de contrôle
    document.add_heading("3. Détail des points de contrôle", level=1)
    _add_tag_paragraph(document, "{% for check in checks %}")

    heading = document.add_paragraph(style="Heading 2")
    _add_tagged_run(heading, "{{ check.rule_id }} — {{ check.title }} — Statut : {{ check.status }}")

    _add_tag_paragraph(document, "{% if check.anomalies %}")

    detail_table = document.add_table(rows=1, cols=6)
    detail_table.style = "Light Grid Accent 1"
    detail_header = detail_table.rows[0].cells
    for index, label in enumerate(
        ["Sévérité", "Description technique", "Preuve technique", "Risque cyber", "Remédiation préconisée", "Référence / Trames"]
    ):
        detail_header[index].paragraphs[0].add_run(label).bold = True
        _set_cell_shading(detail_header[index], "D9E2F3")

    detail_open_marker = detail_table.add_row().cells
    detail_open_marker[0].paragraphs[0].add_run("{%tr for anomaly in check.anomalies %}")

    body_cells = detail_table.add_row().cells
    body_cells[0].paragraphs[0].add_run("{{ anomaly.severity }}")
    body_cells[1].paragraphs[0].add_run("{{ anomaly.description }}")
    body_cells[2].paragraphs[0].add_run("{{ anomaly.technical_evidence }}")
    body_cells[3].paragraphs[0].add_run("{{ anomaly.cyber_risk }}")
    body_cells[4].paragraphs[0].add_run("{{ anomaly.remediation }}")
    body_cells[5].paragraphs[0].add_run(
        "{{ anomaly.source_reference }} (trames : {{ anomaly.frame_numbers }})"
    )

    detail_close_marker = detail_table.add_row().cells
    detail_close_marker[0].paragraphs[0].add_run("{%tr endfor %}")

    _add_tag_paragraph(document, "{% else %}")
    no_anomaly_paragraph = document.add_paragraph()
    no_anomaly_paragraph.add_run("Aucun écart détecté pour ce point de contrôle.").italic = True
    _add_tag_paragraph(document, "{% endif %}")

    document.add_paragraph()
    _add_tag_paragraph(document, "{% endfor %}")

    # ------------------------------------------------------------------ Références
    document.add_heading("4. Références documentaires", level=1)
    document.add_paragraph(
        "Le présent rapport a été produit à partir du référentiel de règles suivant, "
        "compilé depuis la base de connaissances documentaire ILEXIA :"
    )
    _add_tag_paragraph(document, "{%p for source in ruleset_sources %}")
    document.add_paragraph("- {{ source }}")
    _add_tag_paragraph(document, "{%p endfor %}")

    document.add_paragraph()
    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        "Rapport généré automatiquement — IVSA (Ilexia Validation & Security Automator). "
        "Confidentiel — Diffusion restreinte."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = _ILEXIA_GREY

    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(TEMPLATE_PATH))
    return TEMPLATE_PATH


if __name__ == "__main__":
    path = build()
    print(f"Gabarit généré : {path}")
